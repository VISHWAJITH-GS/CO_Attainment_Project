import pandas as pd
from docx import Document
import os
import re
import pytesseract
from PIL import Image
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

# ====================================================
# VS CODE SETUP - DEFINE LOCAL FILE PATHS HERE
# ====================================================
DOCX_FILE = "your_question_paper.docx"  # Replace with your DOCX filename
DB_FILE = "your_student_db.xlsx"        # Replace with your Student DB filename

name, ext = os.path.splitext(DOCX_FILE)

# ---------------- READ DOCX ----------------
def read_docx(path):
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)

    tables = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            tables.append(cells)
    return lines, tables

# ---------------- OCR ----------------
def ocr_from_doc(path):
    doc = Document(path)
    ocr_text = ""
    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            img = doc.part.rels[rel].target_part.blob
            with open("temp.png","wb") as f:
                f.write(img)
            ocr_text += pytesseract.image_to_string(Image.open("temp.png"))
    return ocr_text

# ---------------- SMART FIND ----------------
def find_value_after_keyword(tables, keyword):
    for i, row in enumerate(tables):
        for j, cell in enumerate(row):
            if keyword.lower() in cell.lower():
                if j + 1 < len(row) and row[j+1]:
                    return row[j+1]
                if i + 1 < len(tables):
                    next_row = tables[i+1]
                    if j < len(next_row) and next_row[j]:
                        return next_row[j]
    return "NA"

# ---------------- METADATA ----------------
def extract_metadata(lines, tables, path):
    text = "\n".join(lines)
    ocr_text = ocr_from_doc(path)
    full_text = text + "\n" + ocr_text

    course_code = find_value_after_keyword(tables, "Course Code")
    course_name = find_value_after_keyword(tables, "Course Name")
    faculty     = find_value_after_keyword(tables, "Faculty")

    date = re.search(r"\d{2}\.\d{2}\.\d{4}\(.*?\)", full_text)
    year = re.search(r"Academic year\s*[\d\-]+\s*\(.*?\)", full_text)

    date = date.group() if date else "NA"
    year = year.group() if year else "NA"

    # ---- CO DETAILS ----
    co_data = {}
    for row in tables:
        row_text = " ".join(row)
        co = re.search(r"(CO\d)", row_text)
        tsp = re.search(r"(TPS\d)", row_text)
        mark = re.search(r"\b\d+(?:\.\d+)?\b$", row_text)

        if co:
            key = co.group()
            if key not in co_data:
                co_data[key] = {
                    "tsp": tsp.group() if tsp else "NA",
                    "mark": mark.group() if mark else "NA"
                }

    return course_code, course_name, faculty, year, date, co_data

# ====================================================
# PROCESS DOCX
# ====================================================
def process_docx(path):
    lines, tables = read_docx(path)
    course_code, course_name, faculty, year, date, co_data = extract_metadata(lines, tables, path)

    meta = []
    meta.append(["Course Code", course_code])
    meta.append(["Course Name", course_name])
    meta.append(["Faculty in-Charge", faculty])
    meta.append(["Academic Year", year])
    meta.append(["Date", date])

    for co in co_data:
        meta.append([co, co_data[co]["tsp"], co_data[co]["mark"]])

    meta_df = pd.DataFrame(meta)

    # ----- QUESTIONS -----
    questions, marks, cos = [], [], []
    q_counter = 1
    part = "A"

    for row in tables:
        text = " ".join(row)
        if "Part A" in text: part="A"; q_counter=1
        if "Part B" in text: part="B"; q_counter=1
        if "Part C" in text: part="C"; q_counter=1

        m=c=None
        for r in row:
            if re.fullmatch(r"\d+(?:\.\d+)?", r): m=r
            if re.match(r"CO\d", r): c=r

        if m and c:
            qno = f"{part}{q_counter}" if part in ["A","B"] else text.split()[0]
            questions.append(qno)
            marks.append(m)
            cos.append(c)
            q_counter+=1

    output={}
    for i in range(len(questions)):
        output[i]=[questions[i], marks[i], cos[i]]

    q_df=pd.DataFrame.from_dict(output,orient='index').T

    final=pd.DataFrame()
    final=pd.concat([final, meta_df])
    final=pd.concat([final, pd.DataFrame([[""]])])
    final=pd.concat([final, q_df], ignore_index=True)

    return final

# ====================================================
# RUN QP CONVERSION
# ====================================================
print(f"Processing {DOCX_FILE}...")
result = process_docx(DOCX_FILE)
output_name = name + "_FINAL.xlsx"
result.to_excel(output_name, index=False, header=False)

# ====================================================
# STEP 2 – STUDENT DB INTEGRATION  (FIXED MATRIX)
# ====================================================
print(f"\nProcessing Student DB: {DB_FILE}...")
db = pd.read_excel(DB_FILE, header=None)
regnos, names = [], []

# -------- READ REGNO & NAME --------
for i in range(3, len(db)):
    r = str(db.iloc[i,0]).strip()
    n = str(db.iloc[i,1]).strip()
    if r.lower() in ["nan","none",""]: continue
    if "instruction" in r.lower(): continue
    regnos.append(r)
    names.append("" if n.lower()=="nan" else n)

# -------- READ MARK MATRIX FROM G4 --------
marks_matrix = []
for i in range(3, 3 + len(regnos)):
    row = []
    for j in range(6, len(db.columns)):
        v = db.iloc[i, j]
        row.append("" if pd.isna(v) else v)
    marks_matrix.append(row)

# ====================================================
# WRITE TO OUR QP EXCEL
# ====================================================
wb = load_workbook(output_name)
ws = wb.active

ws["A10"]="S.No"
ws["B10"]="REGNO"
ws["C10"]="Student NAME"

for cell in ["A11","A12","B11","B12","C11","C12"]:
    ws[cell]=""

start = 13

# ----- STUDENT INFO -----
for i in range(len(regnos)):
    ws[f"A{start+i}"]=i+1
    ws[f"B{start+i}"]=regnos[i]
    ws[f"C{start+i}"]=names[i]

# ----- FULL MARK MATRIX -----
for i in range(len(marks_matrix)):
    for j in range(len(marks_matrix[i])):
        ws.cell(row=start+i, column=4+j).value = marks_matrix[i][j]

# ---- AUTO WIDTH ----
for col in ws.columns:
    max_len=0
    col_letter=get_column_letter(col[0].column)
    for cell in col:
        if cell.value:
            max_len=max(max_len,len(str(cell.value)))
    ws.column_dimensions[col_letter].width=max_len+4

wb.save(output_name)
print(f"✅ FULL MATRIX COPIED SUCCESSFULLY → {output_name}")